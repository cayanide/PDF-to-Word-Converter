import pdfplumber
import docx
from pdf2docx import Converter
import re
from docx.shared import RGBColor
import tkinter as tk
from tkinter import filedialog, messagebox
import os

def pdf_to_docx(pdf_path, docx_path):
    """
    Converts a PDF file to a DOCX file while preserving layout.
    :param pdf_path: Path to the input PDF file
    :param docx_path: Path to the output DOCX file
    """
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

def extract_text_from_pdf(pdf_path):
    """
    Extracts raw text from a PDF file.
    :param pdf_path: Path to the input PDF file
    :return: Extracted text as a string
    """
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_links(text):
    """
    Extracts URLs from the given text.
    :param text: Extracted text from the PDF
    :return: List of detected URLs
    """
    url_pattern = r'(https?://\S+|www\.\S+)'
    return re.findall(url_pattern, text)

def format_docx(docx_path):
    """
    Formats the generated DOCX file by aligning text and converting detected URLs into hyperlinks.
    :param docx_path: Path to the DOCX file to be formatted
    """
    doc = docx.Document(docx_path)

    for para in doc.paragraphs:
        para.alignment = 1  # Center align text for better presentation
        words = para.text.split(" ")
        para.clear()

        for word in words:
            if re.match(r'(https?://\S+|www\.\S+)', word):
                add_hyperlink(para, word, word)
            else:
                para.add_run(word + " ")

    doc.save(docx_path)

def add_hyperlink(paragraph, text, url):
    """
    Inserts a hyperlink into a Word document.
    :param paragraph: The paragraph to add the hyperlink to
    :param text: The text to be displayed
    :param url: The hyperlink URL
    """
    r = paragraph.add_run(text)
    r.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
    r.font.underline = True
    r.style = 'Hyperlink'

def browse_files():
    files = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")])
    if files:
        for file in files:
            process_pdf(file)

def process_pdf(pdf_file):
    docx_file = os.path.splitext(pdf_file)[0] + "_converted.docx"
    pdf_to_docx(pdf_file, docx_file)
    format_docx(docx_file)
    messagebox.showinfo("Success", f"Converted and formatted: {docx_file}")

def create_gui():
    root = tk.Tk()
    root.title("PDF to DOCX Converter")
    root.geometry("400x300")

    label = tk.Label(root, text="Select PDFs to Convert", font=("Arial", 12))
    label.pack(pady=10)

    browse_button = tk.Button(root, text="Browse PDF Files", command=browse_files)
    browse_button.pack(pady=10)

    root.mainloop()

if __name__ == "__main__":
    create_gui()
